import os
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.vectorstores.utils import DistanceStrategy
from transformers import AutoTokenizer
from langchain.schema import Document
from docx import Document as DocxDocument

CACHE_DIR = os.path.normpath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "models")
)

class Encoder:
    def __init__(
        self, model_name: str = "sentence-transformers/all-MiniLM-L12-v2", device="cpu"
    ):
        self.embedding_function = HuggingFaceEmbeddings(
            model_name=model_name,
            cache_folder=CACHE_DIR,
            model_kwargs={"device": device},
        )

class FaissDb:
    def __init__(self, docs, embedding_function):
        self.db = FAISS.from_documents(
            docs, embedding_function, distance_strategy=DistanceStrategy.COSINE
        )

    def similarity_search(self, question: str, k: int = 3):
        retrieved_docs = self.db.similarity_search(question, k=k)
        context = "".join(doc.page_content + "\n" for doc in retrieved_docs)
        return context
def clean_text(text: str) -> str:
    """
    Làm sạch văn bản: Loại bỏ các ký tự xuống dòng (\n) và khoảng trắng thừa.
    """
    return " ".join(text.split())

def load_and_split_files(file_paths: list, chunk_size: int = 256):
    loaders = []
    pages = []

    for file_path in file_paths:
        if file_path.endswith(".pdf"):
            loaders.append(PyPDFLoader(file_path))
        elif file_path.endswith(".docx"):
            try:
                doc = DocxDocument(file_path)
                text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # Chỉ thêm đoạn có nội dung
                        text.append(clean_text(paragraph.text))
                if text:
                    # Tạo đối tượng Document từ langchain.schema
                    pages.append(Document(page_content=" ".join(text), metadata={"source": file_path}))
                else:
                    print(f"No content found in DOCX file: {file_path}")
            except Exception as e:
                print(f"Error loading DOCX {file_path}: {e}")
        else:
            print(f"Unsupported file format: {file_path}")

    # Xử lý file PDF
    for loader in loaders:
        try:
            loaded_pages = loader.load()
            for page in loaded_pages:
                # Làm sạch văn bản và chuyển thành đối tượng Document
                page_content = clean_text(page.page_content)
                pages.append(Document(page_content=page_content, metadata={"source": loader.file_path}))
            print(f"Loaded {len(loaded_pages)} pages from {loader.file_path}")
        except Exception as e:
            print(f"Error loading PDF {loader.file_path}: {e}")

    if not pages:
        print("No pages loaded from the provided files.")
        return []

    # Tách đoạn văn bản với RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        tokenizer=AutoTokenizer.from_pretrained(
            "sentence-transformers/all-MiniLM-L12-v2"
        ),
        chunk_size=chunk_size,
        chunk_overlap=int(chunk_size / 10),
        strip_whitespace=True,
    )
    docs = text_splitter.split_documents(pages)
    return docs